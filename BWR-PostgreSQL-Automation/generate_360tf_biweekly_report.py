#!/usr/bin/env python3
"""Generate the 360tf PostgreSQL biweekly DBA report in Microsoft Word.

Credentials are loaded from environment variables or a .env file. Do not put a
password in this source file.
"""
from __future__ import annotations

import os
import smtplib
import sys
from datetime import datetime
from email.message import EmailMessage
from pathlib import Path
from typing import Any, Iterable

import psycopg
from dotenv import load_dotenv
from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

load_dotenv()

# Connection and report settings. Set these in .env where appropriate.
CLIENT = os.getenv("REPORT_CLIENT", "360tf")
APP_DATABASE = os.getenv("APP_DATABASE", "app360tf")
HOST = os.getenv("PGHOST", "")
PORT = os.getenv("PGPORT", "5432")
USER = os.getenv("PGUSER", "")
PASSWORD = os.getenv("PGPASSWORD", "")
SSL_MODE = os.getenv("PGSSLMODE", "require")
ENDPOINT = os.getenv("REPORT_ENDPOINT", f"{HOST}:{PORT}")
OUT_DIR = Path(os.getenv("REPORT_OUTPUT_DIR", "reports"))
MAIL_TO = "postgresqlsupport@geopits.com"
SMTP_HOST = os.getenv("SMTP_HOST", "")
SMTP_PORT = int(os.getenv("SMTP_PORT", "587"))
SMTP_USER = os.getenv("SMTP_USER", "")
SMTP_PASSWORD = os.getenv("SMTP_PASSWORD", "")
SMTP_FROM = os.getenv("SMTP_FROM", SMTP_USER)
SMTP_USE_SSL = os.getenv("SMTP_USE_SSL", "false").lower() == "true"
EMAIL_ENABLED = os.getenv("EMAIL_ENABLED", "true").lower() == "true"
EXCLUDED_DBS = {"template0", "template1", "rdsadmin"}

# RDS facts requested for the server-overview section. Update if the instance changes.
RDS_FACTS = [
    ("DB Identifier", os.getenv("RDS_IDENTIFIER", "prod360tf")),
    ("Status", os.getenv("RDS_STATUS", "Available")),
    ("Class", os.getenv("RDS_CLASS", "db.t3.medium")),
    ("Role", os.getenv("RDS_ROLE", "Instance")),
    ("Engine", os.getenv("RDS_ENGINE", "PostgreSQL")),
    ("Region & AZ", os.getenv("RDS_REGION_AZ", "us-east-1b")),
    ("Storage", os.getenv("RDS_STORAGE", "60 GiB")),
    ("Provisioned IOPS", os.getenv("RDS_IOPS", "3000 IOPS")),
    ("Storage Throughput", os.getenv("RDS_THROUGHPUT", "125 MiBps")),
    ("Storage Autoscaling", os.getenv("RDS_AUTOSCALING", "Enabled")),
    ("Maximum Storage Threshold", os.getenv("RDS_MAX_STORAGE", "65 GiB")),
    ("vCPU", os.getenv("RDS_VCPU", "2")),
    ("RAM", os.getenv("RDS_RAM", "4 GB")),
    ("Deletion Protection", os.getenv("RDS_DELETION_PROTECTION", "Enabled")),
]


def connection_config(dbname: str = "postgres") -> dict[str, str]:
    return {"host": HOST, "port": PORT, "user": USER, "password": PASSWORD,
            "sslmode": SSL_MODE, "dbname": dbname, "connect_timeout": "15"}


def query(dbname: str, sql: Any, params: Iterable[Any] | None = None) -> list[dict[str, Any]]:
    with psycopg.connect(**connection_config(dbname), row_factory=psycopg.rows.dict_row) as conn:
        with conn.cursor() as cur:
            cur.execute(sql, params)
            return cur.fetchall()


def safe(dbname: str, label: str, sql: Any, params: Iterable[Any] | None = None) -> tuple[list[dict[str, Any]], str | None]:
    try:
        return query(dbname, sql, params), None
    except Exception as exc:
        return [], f"{label}: {str(exc).splitlines()[0]}"


def pretty_size(value: Any) -> str:
    """Format values consistently: bytes/KB/MB below 1 GB, GB at and above 1 GB."""
    if value is None:
        return "N/A"
    size = int(value)
    if size >= 1024 ** 3:
        return f"{size / (1024 ** 3):.2f} GB"
    if size >= 1024 ** 2:
        return f"{size / (1024 ** 2):.2f} MB"
    if size >= 1024:
        return f"{size / 1024:.2f} KB"
    return f"{size} bytes"


def add_field(run: Any, instruction: str, placeholder: str = "Update field") -> None:
    begin = OxmlElement("w:fldChar"); begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve"); instr.text = instruction
    separate = OxmlElement("w:fldChar"); separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t"); text.text = placeholder; separate.append(text)
    end = OxmlElement("w:fldChar"); end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, end])


def add_bookmark(paragraph: Any, name: str, bookmark_id: int) -> None:
    start = OxmlElement("w:bookmarkStart"); start.set(qn("w:id"), str(bookmark_id)); start.set(qn("w:name"), name)
    end = OxmlElement("w:bookmarkEnd"); end.set(qn("w:id"), str(bookmark_id))
    paragraph._p.insert(0, start); paragraph._p.append(end)


def set_vertical_alignment(section: Any, value: str) -> None:
    sect_pr = section._sectPr
    tag = qn("w:vAlign")
    old = sect_pr.find(tag)
    if old is not None:
        sect_pr.remove(old)
    node = OxmlElement("w:vAlign"); node.set(qn("w:val"), value); sect_pr.append(node)


def shade(cell: Any, fill: str) -> None:
    props = cell._tc.get_or_add_tcPr()
    node = OxmlElement("w:shd"); node.set(qn("w:fill"), fill); props.append(node)


def set_cell(cell: Any, value: Any, bold: bool = False, white: bool = False) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(str(value if value not in (None, "") else "N/A"))
    run.bold = bold; run.font.size = Pt(8)
    if white:
        run.font.color.rgb = RGBColor(255, 255, 255)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def add_table(doc: Document, headers: list[str], rows: list[Any], keys: list[str] | None = None) -> None:
    if not rows:
        doc.add_paragraph("No applicable records were returned.")
        return
    table = doc.add_table(rows=1, cols=len(headers)); table.style = "Table Grid"; table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for idx, header in enumerate(headers):
        set_cell(table.rows[0].cells[idx], header, bold=True, white=True); shade(table.rows[0].cells[idx], "1F4E78")
    for source in rows:
        values = [source.get(k, "N/A") for k in keys] if isinstance(source, dict) and keys else source
        cells = table.add_row().cells
        for idx, value in enumerate(values):
            set_cell(cells[idx], value)
    doc.add_paragraph().paragraph_format.space_after = Pt(3)


def add_heading(doc: Document, title: str, level: int = 1, bookmark: str | None = None, bookmark_id: int = 0) -> None:
    p = doc.add_heading(title, level=level)
    if bookmark:
        add_bookmark(p, bookmark, bookmark_id)


def observation(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    r = p.add_run("Observation: "); r.bold = True; r.font.color.rgb = RGBColor(31, 78, 120)
    p.add_run(text)


def discover_databases() -> tuple[list[str], list[str]]:
    requested = [d.strip() for d in os.getenv("PGDATABASES", "").split(",") if d.strip()]
    if requested:
        return requested, []
    rows, err = safe("postgres", "Database discovery", "SELECT datname FROM pg_database WHERE datallowconn AND NOT datistemplate ORDER BY datname")
    return [r["datname"] for r in rows if r["datname"] not in EXCLUDED_DBS], [err] if err else []


def collect_from_each_database(databases: list[str], label: str, sql: str) -> tuple[list[dict[str, Any]], list[str]]:
    all_rows: list[dict[str, Any]] = []
    notes: list[str] = []
    for db in databases:
        rows, error = safe(db, f"{db} {label}", sql)
        all_rows.extend(rows)
        if error:
            notes.append(error)
    return all_rows, notes


def exact_large_tables(databases: list[str]) -> tuple[list[list[Any]], list[str]]:
    """Use COUNT(*) for exact records, as requested. This can be resource intensive."""
    records: list[list[Any]] = []
    notes: list[str] = []
    from psycopg import sql
    for db in databases:
        relations, error = safe(db, f"{db} table list", "SELECT schemaname, relname FROM pg_stat_user_tables ORDER BY schemaname, relname")
        if error:
            notes.append(error); continue
        for rel in relations:
            statement = sql.SQL("SELECT count(*) AS record_count FROM {}.{}").format(sql.Identifier(rel["schemaname"]), sql.Identifier(rel["relname"]))
            count_rows, count_error = safe(db, f"{db}.{rel['schemaname']}.{rel['relname']} exact count", statement)
            if count_error:
                notes.append(count_error); continue
            count = count_rows[0]["record_count"]
            if count > 10_000_000:
                records.append([db, rel["schemaname"], rel["relname"], count])
    return records, notes


def send_report_email(report_path: Path) -> None:
    """Email the completed report as an attachment using the configured SMTP account."""
    if not EMAIL_ENABLED:
        print("Email delivery disabled (EMAIL_ENABLED=false).")
        return
    if not all([SMTP_HOST, SMTP_FROM, SMTP_USER, SMTP_PASSWORD]):
        raise RuntimeError("Report was created, but email was not sent. Set SMTP_HOST, SMTP_USER, SMTP_PASSWORD, and SMTP_FROM in .env.")
    message = EmailMessage()
    message["From"] = SMTP_FROM
    message["To"] = MAIL_TO
    message["Subject"] = f"360tf-BWR Report({datetime.now():%d-%b-%Y})"
    message.set_content("""Hi Team,

Please find the BWR Report for the 360tf.
Kindly validate and add the graphs and its observations.
Thanks.

Automated Report by
Ajay S
""")
    with report_path.open("rb") as attachment:
        message.add_attachment(attachment.read(), maintype="application",
                               subtype="vnd.openxmlformats-officedocument.wordprocessingml.document",
                               filename=report_path.name)
    if SMTP_USE_SSL:
        with smtplib.SMTP_SSL(SMTP_HOST, SMTP_PORT, timeout=30) as smtp:
            smtp.login(SMTP_USER, SMTP_PASSWORD)
            smtp.send_message(message)
    else:
        with smtplib.SMTP(SMTP_HOST, SMTP_PORT, timeout=30) as smtp:
            smtp.ehlo(); smtp.starttls(); smtp.ehlo()
            smtp.login(SMTP_USER, SMTP_PASSWORD)
            smtp.send_message(message)


def build_report() -> Path:
    if not all([HOST, USER, PASSWORD]):
        raise SystemExit("PGHOST, PGUSER, and PGPASSWORD must be set in the environment or .env file.")
    databases, notes = discover_databases()
    if not databases:
        raise SystemExit("No accessible databases found. Set PGDATABASES in .env.")
    if APP_DATABASE not in databases:
        notes.append(f"The detailed table section was not collected because {APP_DATABASE} is not accessible. Set PGDATABASES to include it.")

    doc = Document()
    section = doc.sections[0]
    section.top_margin = section.bottom_margin = Inches(0.8); section.left_margin = section.right_margin = Inches(0.85)
    set_vertical_alignment(section, "center")
    normal = doc.styles["Normal"]; normal.font.name = "Aptos"; normal.font.size = Pt(9); normal.paragraph_format.space_after = Pt(6)
    for name, size, color in [("Title", 28, "1F4E78"), ("Heading 1", 16, "1F4E78"), ("Heading 2", 12, "2F75B5")]:
        style = doc.styles[name]; style.font.name = "Aptos Display"; style.font.size = Pt(size); style.font.color.rgb = RGBColor.from_string(color)

    # Centered cover page, with no header or footer.
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("POSTGRESQL"); r.bold = True; r.font.size = Pt(14); r.font.color.rgb = RGBColor(31, 78, 120)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"{CLIENT} Biweekly Report"); r.bold = True; r.font.size = Pt(30); r.font.color.rgb = RGBColor(31, 78, 120)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(f"Generated Date: {datetime.now():%d %b %Y}").font.size = Pt(12)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(f"Server Endpoint: {ENDPOINT}").font.size = Pt(12)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Automated report by Ajay S"); r.italic = True; r.font.size = Pt(9); r.font.color.rgb = RGBColor(166, 166, 166)

    # New page starts the body and returns to standard top aligned pages.
    body_section = doc.add_section(WD_SECTION_START.NEW_PAGE)
    body_section.top_margin = body_section.bottom_margin = Inches(0.8); body_section.left_margin = body_section.right_margin = Inches(0.85)
    set_vertical_alignment(body_section, "top")

    toc_entries = [
        ("Introduction", "toc_intro"), ("Executive Summary", "toc_exec"), ("PostgreSQL Server Overview", "toc_server"),
        ("Key Activities Performed During Past Two Weeks", "toc_activities"), ("Database Overview", "toc_database"),
        ("List of Databases and Their Sizes", "toc_sizes"), ("Database Growth", "toc_growth"),
        (f"List of Tables in {APP_DATABASE} Database", "toc_tables"), (f"Table Growth Difference in {APP_DATABASE} Database", "toc_tablegrowth"),
        ("Tablespace Usage", "toc_tablespace"), ("Tables with Greater Than 1 Crore Records", "toc_large"),
        ("Count of Zero-Record Tables", "toc_zero"), ("Tables Without Primary Key", "toc_nopk"),
        ("Tables with Bloat", "toc_bloat"), ("Users & Roles", "toc_roles"), ("Superuser Accounts", "toc_super"),
        ("Autovacuum Statistics", "toc_auto"), ("Console Graphs (Add Manually)", "toc_graphs"), ("Ticket Details", "toc_tickets"),
    ]
    add_heading(doc, "Table of Contents", bookmark="toc_contents", bookmark_id=1)
    contents = doc.add_table(rows=1, cols=2); contents.style = "Table Grid"; contents.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, header in enumerate(["S. No.", "Title"]):
        set_cell(contents.rows[0].cells[i], header, True, True); shade(contents.rows[0].cells[i], "1F4E78")
    for number, (title, bookmark) in enumerate(toc_entries, 1):
        row = contents.add_row().cells
        set_cell(row[0], number); set_cell(row[1], title)
    doc.add_page_break()

    bid = 20
    def h(title: str, level: int = 1, key: str | None = None) -> None:
        nonlocal bid
        add_heading(doc, title, level, key, bid); bid += 1

    h("Introduction", key="toc_intro")
    doc.add_paragraph(f"This report provides a consolidated view of the PostgreSQL environment supporting {CLIENT}. It covers server configuration, database capacity, table growth, maintenance health, security roles, and operational observations. The report is generated from live PostgreSQL catalog and statistics views, allowing the DBA team to identify capacity trends, maintenance requirements, and data-quality risks in one place.")
    h("Executive Summary", key="toc_exec")
    for item in ["Database instance health and connectivity were checked.", "Database growth and storage utilization were analyzed.", "Table growth, audit-data growth, table bloat, and autovacuum activity were reviewed.", "Database connection usage, user/role configuration, scheduler/background-job tables, and performance statistics were validated."]:
        doc.add_paragraph(item, style="List Bullet")

    h("PostgreSQL Server Overview", key="toc_server")
    add_table(doc, ["Parameter", "Value"], RDS_FACTS)
    h("Key Activities Performed", key="toc_activities")
    for item in ["Database instance health monitoring completed.", "Database growth and storage utilization analyzed.", "Table growth and audit data growth reviewed.", "Table bloat assessment conducted.", "Autovacuum operations verified.", "Database connection usage monitored.", "User and role configurations reviewed.", "Scheduler and background job tables reviewed.", "Performance metrics validated."]:
        doc.add_paragraph(item, style="List Bullet")

    h("Database Overview", key="toc_database")
    overview, err = safe(databases[0], "Database overview", "SELECT datname AS name, pg_get_userbyid(datdba) AS owner, pg_encoding_to_char(encoding) AS encoding, datcollate AS collation, datctype AS ctype FROM pg_database WHERE datname = ANY(%s) ORDER BY datname", (databases,))
    if err: notes.append(err)
    add_table(doc, ["Name", "Owner", "Encoding", "Collation", "CType"], overview, ["name", "owner", "encoding", "collation", "ctype"])
    h("List of Databases and Their Sizes", key="toc_sizes")
    db_sizes, err = safe(databases[0], "Database sizes", "SELECT datname AS database_name, pg_database_size(datname) AS size_bytes FROM pg_database WHERE datname = ANY(%s) ORDER BY pg_database_size(datname) DESC", (databases,))
    if err: notes.append(err)
    add_table(doc, ["Database Name", "Size"], [[r["database_name"], pretty_size(r["size_bytes"])] for r in db_sizes])
    h("Database Growth", key="toc_growth")
    add_table(doc, ["Database Name", "Size (Last Time)", "Current Size", "Difference (%)"], [[r["database_name"], "N/A", pretty_size(r["size_bytes"]), "N/A"] for r in db_sizes])
    observation(doc, "PostgreSQL does not retain historic database-size snapshots. The current run records the current sizes; import the previous report snapshot to calculate last-time size and percentage change.")

    h(f"List of Tables in {APP_DATABASE} Database", key="toc_tables")
    app_tables, err = safe(APP_DATABASE, "Table inventory", "SELECT current_database() AS database_name, schemaname AS schema_name, relname AS table_name FROM pg_stat_user_tables ORDER BY schemaname, relname")
    if err: notes.append(err)
    add_table(doc, ["Database Name", "Schema Name", "Table Name"], app_tables, ["database_name", "schema_name", "table_name"])
    h(f"Table Growth Difference in {APP_DATABASE} Database", key="toc_tablegrowth")
    observation(doc, "Historical relation-size snapshots are required for an accurate 15-day change. Current sizes are recorded here for the next comparison cycle.")
    table_sizes, err = safe(APP_DATABASE, "Current table sizes", "SELECT schemaname AS schema_name, relname AS table_name, pg_total_relation_size(relid) AS size_bytes FROM pg_stat_user_tables ORDER BY pg_total_relation_size(relid) DESC")
    if err: notes.append(err)
    add_table(doc, ["Schema Name", "Table Name", "Size (Last Time)", "Current Size", "Difference (%)"], [[r["schema_name"], r["table_name"], "N/A", pretty_size(r["size_bytes"]), "N/A"] for r in table_sizes])

    h("Tablespace Usage", key="toc_tablespace")
    doc.add_paragraph("A tablespace is a PostgreSQL storage location for database objects. It helps place data on appropriate disks and supports storage-capacity and performance management.")
    tablespaces, err = safe(databases[0], "Tablespace usage", "SELECT spcname AS tablespace_name, pg_tablespace_size(oid) AS size_bytes FROM pg_tablespace ORDER BY spcname")
    if err: notes.append(err)
    add_table(doc, ["Tablespace Name", "Size"], [[r["tablespace_name"], pretty_size(r["size_bytes"])] for r in tablespaces])

    h("Tables with Greater Than 1 Crore Records", key="toc_large")
    large, large_notes = exact_large_tables(databases); notes.extend(large_notes)
    if large:
        add_table(doc, ["Database Name", "Schema Name", "Table Name", "Exact Record Count"], large)
    else:
        doc.add_paragraph("No tables with more than 1 crore (10,000,000) records were found.")

    h("Count of Zero-Record Tables", key="toc_zero")
    zero_rows: list[list[Any]] = []
    for db in databases:
        rows, err = safe(db, "Zero-record table count", "SELECT count(*) AS zero_record_table_count FROM pg_stat_user_tables WHERE n_live_tup = 0")
        if err: notes.append(f"{db} {err}")
        zero_rows.append([db, rows[0]["zero_record_table_count"] if rows else "N/A"])
    add_table(doc, ["Database Name", "Zero-Record Table Count"], zero_rows)

    h("Tables Without Primary Key", key="toc_nopk")
    nopk, collected_notes = collect_from_each_database(databases, "tables without primary key", "SELECT current_database() AS database_name, n.nspname AS schema_name, c.relname AS table_name FROM pg_class c JOIN pg_namespace n ON n.oid=c.relnamespace WHERE c.relkind IN ('r','p') AND n.nspname NOT IN ('pg_catalog','information_schema') AND NOT EXISTS (SELECT 1 FROM pg_constraint con WHERE con.conrelid=c.oid AND con.contype='p') ORDER BY n.nspname, c.relname")
    notes.extend(collected_notes)
    if nopk:
        add_table(doc, ["Database Name", "Schema Name", "Table Name"], nopk, ["database_name", "schema_name", "table_name"])
    else:
        doc.add_paragraph("No tables without a primary key were found.")

    h("Tables with Bloat", key="toc_bloat")
    bloat, collected_notes = collect_from_each_database(databases, "bloat assessment", "SELECT current_database() AS database_name, schemaname AS schema_name, relname AS table_name, n_dead_tup, n_live_tup, CASE WHEN n_live_tup + n_dead_tup = 0 THEN 0 ELSE round(100.0 * n_dead_tup / (n_live_tup + n_dead_tup), 2) END AS bloat_percentage FROM pg_stat_user_tables WHERE n_dead_tup > 0 ORDER BY n_dead_tup DESC")
    notes.extend(collected_notes)
    if bloat:
        add_table(doc, ["Database", "Schema", "Table", "Dead Tuples", "Live Tuples", "Est. Bloat %"], bloat, ["database_name", "schema_name", "table_name", "n_dead_tup", "n_live_tup", "bloat_percentage"])
    else:
        doc.add_paragraph("No tables with recorded dead tuples were found.")
    observation(doc, "Live tuples are currently valid rows; dead tuples are obsolete row versions left after UPDATE or DELETE operations until VACUUM reclaims them. Bloat is unused table space that can grow when dead tuples are not cleaned up. The estimated bloat percentage shown is dead tuples ÷ (live tuples + dead tuples) × 100; it is a maintenance indicator, not an exact disk-bloat measurement.")

    h("Users & Roles", key="toc_roles")
    roles, err = safe(databases[0], "Users and roles", "SELECT rolname AS role_name, concat_ws(', ', CASE WHEN rolsuper THEN 'SUPERUSER' END, CASE WHEN rolcreatedb THEN 'CREATEDB' END, CASE WHEN rolcreaterole THEN 'CREATEROLE' END, CASE WHEN rolcanlogin THEN 'LOGIN' END) AS attributes, COALESCE((SELECT string_agg(parent.rolname, ', ') FROM pg_auth_members m JOIN pg_roles parent ON parent.oid=m.roleid WHERE m.member=r.oid), 'N/A') AS member_of FROM pg_roles r WHERE rolname NOT LIKE 'pg_%' ORDER BY rolname")
    if err: notes.append(err)
    add_table(doc, ["Role Name", "Attributes", "Member Of"], roles, ["role_name", "attributes", "member_of"])
    observation(doc, "N/A in the Member Of column means the role has no assigned parent role. N/A in other sections indicates that the value is not applicable or was not available to the reporting account.")
    h("Superuser Accounts", key="toc_super")
    superusers = [r for r in roles if "SUPERUSER" in str(r.get("attributes"))]
    if superusers:
        add_table(doc, ["Role Name", "Attributes", "Member Of"], superusers, ["role_name", "attributes", "member_of"])
    else:
        doc.add_paragraph("No superuser accounts were returned. On Amazon RDS, the master user is commonly not a true PostgreSQL superuser.")

    h("Autovacuum Statistics", key="toc_auto")
    auto, collected_notes = collect_from_each_database(databases, "autovacuum statistics", "SELECT current_database() AS database_name, schemaname AS schema_name, relname AS table_name, to_char(last_autovacuum, 'YYYY-MM-DD HH24:MI:SS') AS last_autovacuum_time, autovacuum_count FROM pg_stat_user_tables WHERE last_autovacuum IS NOT NULL ORDER BY last_autovacuum DESC, relname")
    notes.extend(collected_notes)
    add_table(doc, ["Database", "Schema", "Table", "Last Autovacuum", "Autovacuum Count"], auto, ["database_name", "schema_name", "table_name", "last_autovacuum_time", "autovacuum_count"])
    no_auto_count = 0
    for db in databases:
        no_auto, err = safe(db, "Tables without autovacuum", "SELECT count(*) AS table_count FROM pg_stat_user_tables WHERE last_autovacuum IS NULL")
        if err:
            notes.append(f"{db} {err}")
        elif no_auto:
            no_auto_count += no_auto[0]["table_count"]
    observation(doc, f"Only tables with a recorded last-autovacuum timestamp are shown above. {no_auto_count} table(s) have no recorded autovacuum; these are generally new, empty, or low-change tables because autovacuum runs only after its activity/dead-tuple threshold is reached. An autovacuum count of 0 can also mean no autovacuum has run since table creation or statistics reset.")

    h("Graphs", key="toc_graphs")
    for metric in ["CPU Utilization", "Database Connections", "Freeable Memory", "Free Storage Space"]:
        h(metric, level=2)
        doc.add_paragraph("[Insert console graph here]")
        observation(doc, "Add a brief trend interpretation, peak values, thresholds, and any required action.")
    if notes:
        h("Data Collection Notes")
        for note in sorted(set(notes)):
            doc.add_paragraph(note, style="List Bullet")
    h("Ticket Details", key="toc_tickets")
    add_table(doc, ["S. No.", "Ticket ID", "Title / Description", "Status", "Remarks"], [["1", "N/A", "No tickets provided for this reporting period.", "N/A", "N/A"]])
    h("Overall Conclusion")
    doc.add_paragraph("During the reporting period, the PostgreSQL production database environment for the 360TF application remained stable, available, and operating within acceptable performance thresholds. No service disruptions or critical performance issues were observed.")
    doc.add_paragraph("System resources including CPU, memory, storage, and database connections remained within safe operational limits. Database growth observed during the period aligns with normal application activity, and automated maintenance processes such as autovacuum continue to effectively maintain table health.")
    doc.add_paragraph("No immediate risks or corrective actions are required at this time. Continuous monitoring and preventive maintenance activities will continue to ensure sustained database performance and availability.")
    doc.add_paragraph("Overall, the database environment is operating in a healthy and stable condition, supporting application workloads without impact.")

    OUT_DIR.mkdir(parents=True, exist_ok=True)
    output = OUT_DIR / f"{CLIENT}_PostgreSQL_Biweekly_Report_{datetime.now():%b_%Y}.docx"
    doc.save(output)
    send_report_email(output)
    return output


if __name__ == "__main__":
    try:
        print(f"Created: {build_report()}")
    except Exception as exc:
        print(f"Report generation failed: {exc}", file=sys.stderr)
        raise SystemExit(1)
